# ollama_client.py
# Use Ollama (local or cloud) to generate an AI summary for word data; output as docx report.
# Loads ollama.env and word.env; structured summary for LLM, system + user prompts,
# Ollama Cloud (chat API) vs local (generate API), error handling.

import os
import json
import requests
from pathlib import Path
from io import BytesIO

# Load ollama.env, word.env, .env for OLLAMA_API_KEY and OLLAMA_MODEL
_here = Path(__file__).resolve().parent
for parent in [_here, _here.parent, _here.parent.parent]:
    for name in ("ollama.env", "word.env", ".env"):
        env_file = parent / name
        if env_file.exists():
            from dotenv import load_dotenv
            load_dotenv(env_file)
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

OLLAMA_API_KEY = os.getenv("OLLAMA_API_KEY")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:latest")
OLLAMA_TIMEOUT = int(os.getenv("OLLAMA_TIMEOUT", "300"))
OLLAMA_PORT = 11434


def _build_summary_for_llm(word_data):
    """Build a short summary for the LLM (fewer tokens = faster, less timeout)."""
    if not word_data or not word_data.get("ok"):
        return json.dumps({"error": word_data.get("error", "No data")})
    word = word_data.get("word", "")
    prons = word_data.get("pronunciations", [])
    defs = word_data.get("definitions", [])
    audio = word_data.get("audio", [])
    summary = {
        "word": word,
        "pronunciations": [{"raw": p.get("raw"), "type": p.get("rawType")} for p in prons],
        "definitions": [{"partOfSpeech": d.get("partOfSpeech"), "text": (d.get("text") or "")[:150]} for d in defs],
        "has_audio": len(audio) > 0,
    }
    return json.dumps(summary, indent=2)


# --- Call Ollama (local or cloud) to get narrative summary ---
SYSTEM_PROMPT = """You are a language coach and data analyst. Given dictionary data for a word (pronunciations and definitions), write a brief report using bullet points. Include:
- How to pronounce the word (use the phonetic spellings provided)
- One or two main meanings
- A simple tip for remembering the pronunciation or using the word
- One follow-up suggestion (e.g., a similar word to learn or a usage example)
Be concise; use bullet points for easy reading."""


def _call_ollama(summary_text):
    """
    Call Ollama (Cloud if OLLAMA_API_KEY set, else local) and return (output_text, error_message).
    Calls Ollama Cloud (chat API) or local (generate API).
    """
    user_prompt = f"Word data (JSON):\n\n{summary_text}\n\nWrite a brief pronunciation and usage report in bullet points."

    if OLLAMA_API_KEY:
        url = "https://ollama.com/api/chat"
        headers = {
            "Authorization": f"Bearer {OLLAMA_API_KEY}",
            "Content-Type": "application/json",
        }
        body = {
            "model": "gpt-oss:20b-cloud",
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            "stream": False,
        }
        try:
            resp = requests.post(url, headers=headers, json=body, timeout=120)
            resp.raise_for_status()
            out = resp.json().get("message", {}).get("content")
            return (out, None) if out else (None, "Ollama Cloud returned no content.")
        except requests.RequestException as e:
            err_msg = str(e)
            if getattr(e, "response", None) and getattr(e.response, "text", None):
                err_msg += " " + (e.response.text or "")[:200]
            return None, f"Ollama Cloud request failed: {err_msg}"

    # Local Ollama (generate API)
    url = f"http://localhost:{OLLAMA_PORT}/api/generate"
    body = {
        "model": OLLAMA_MODEL,
        "prompt": f"{SYSTEM_PROMPT}\n\n---\n\n{user_prompt}",
        "stream": False,
    }
    try:
        resp = requests.post(url, json=body, timeout=OLLAMA_TIMEOUT)
        resp.raise_for_status()
        out = resp.json().get("response", "")
        return (out.strip() or None, None) if out else (None, "Ollama returned empty response.")
    except requests.ConnectionError as e:
        return None, "Ollama is not running. Start it with: ollama serve — then run: ollama pull " + (OLLAMA_MODEL or "gemma3:latest")
    except requests.Timeout as e:
        return None, "Ollama timed out. Try a smaller model (e.g. OLLAMA_MODEL=llama3.2:3b) or set OLLAMA_TIMEOUT=600 in ollama.env."
    except requests.RequestException as e:
        err_detail = ""
        if getattr(e, "response", None):
            try:
                j = e.response.json()
                if "error" in j:
                    err_detail = " " + str(j["error"])
            except Exception:
                if getattr(e.response, "text", None):
                    err_detail = " " + (e.response.text or "")[:200]
        if "404" in str(e):
            err_detail = " Model not found — try: ollama pull " + (OLLAMA_MODEL or "gemma3:latest") + err_detail
        return None, f"Ollama request failed:{err_detail}"


def generate_report_docx(word_data):
    """
    Generate a Word document (docx) containing word data and an AI summary from Ollama.
    Structured summary from word data -> Ollama -> narrative in report.

    Parameters
    ----------
    word_data : dict
        Result from api_client.get_word_data() (word, pronunciations, definitions, audio).

    Returns
    -------
    bytes
        The .docx file content.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()
    doc.add_heading("Word Pronunciation Report", 0)

    if not word_data or not word_data.get("ok"):
        doc.add_paragraph(word_data.get("error", "No word data.") if word_data else "No word data.")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    word = word_data.get("word", "")
    prons = word_data.get("pronunciations", [])
    defs = word_data.get("definitions", [])
    fetched_at = word_data.get("fetched_at", "")

    doc.add_heading(word, level=1)
    if fetched_at:
        doc.add_paragraph(f"Fetched: {fetched_at}", style="Intense Quote")

    doc.add_heading("Pronunciations", level=2)
    if prons:
        for p in prons:
            raw = p.get("raw") or "—"
            raw_type = p.get("rawType") or ""
            doc.add_paragraph(f"{raw} ({raw_type})")
    else:
        doc.add_paragraph("No pronunciations found.")

    doc.add_heading("Definitions", level=2)
    if defs:
        for d in defs:
            pos = d.get("partOfSpeech") or ""
            text = (d.get("text") or "").strip()
            doc.add_paragraph(f"[{pos}] {text}")
    else:
        doc.add_paragraph("No definitions found.")

    doc.add_heading("AI Summary", level=2)
    summary_text = _build_summary_for_llm(word_data)
    output, err = _call_ollama(summary_text)
    if err:
        doc.add_paragraph("AI summary could not be generated:")
        doc.add_paragraph(err, style="Intense Quote")
    else:
        doc.add_paragraph(output or "")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# -----------------------------------------------------------------------------
# Save AI report in multiple formats (.txt, .md, .html, .docx) — like 05_reporting
# -----------------------------------------------------------------------------

def get_report_text(word_data):
    """
    Get the AI report narrative only (no word data). Returns (text, error).
    Use this to then save in .txt, .md, .html, or .docx via the helpers below.
    """
    if not word_data or not word_data.get("ok"):
        return None, (word_data.get("error") if word_data else "No word data.") or "No data"
    summary_text = _build_summary_for_llm(word_data)
    return _call_ollama(summary_text)


def report_to_txt(report_text):
    """Plain text format. Returns str."""
    return report_text or ""


def report_to_md(report_text):
    """Markdown format (same content). Returns str."""
    return report_text or ""


def report_to_html(report_text, title="Word Pronunciation AI Report"):
    """Convert report to HTML (markdown -> HTML with wrapper). Returns str."""
    if not report_text:
        return f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{title}</title></head><body><p>No report content.</p></body></html>"
    try:
        import markdown
        html_content = markdown.markdown(report_text)
    except Exception:
        html_content = "<pre>" + (report_text or "").replace("<", "&lt;").replace(">", "&gt;") + "</pre>"
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ color: #333; }}
        h2 {{ color: #666; margin-top: 30px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""


def report_to_docx_ai_only(report_text):
    """
    Build a Word document containing only the AI report text.
    Parses markdown-style lines: # heading 1, ## heading 2, - bullet, else paragraph.
    Returns bytes.
    """
    from docx import Document
    doc = Document()
    if not report_text:
        doc.add_paragraph("No report content.")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    for line in report_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# -----------------------------------------------------------------------------
# Run from command line to produce a docx file
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    import sys
    from api_client import get_word_data

    word = (sys.argv[1].strip() if len(sys.argv) > 1 else "").strip() or "hello"
    print(f"Fetching data for '{word}'...")
    word_data = get_word_data(word)
    if not word_data.get("ok"):
        print("Error:", word_data.get("error", "No data"))
        sys.exit(1)
    word_data["fetched_at"] = __import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M")
    print("Generating report (calling Ollama)...")
    docx_bytes = generate_report_docx(word_data)
    out_name = f"word_report_{word}.docx"
    Path(out_name).write_bytes(docx_bytes)
    print(f"Saved: {out_name}")
